"""CV Parser - Supports PDF, DOCX, TXT, JSON formats"""

import json
import os
from typing import Dict, Optional
import logging

logger = logging.getLogger(__name__)


class CVParser:
    """Parse CVs from multiple formats"""
    
    @staticmethod
    def parse_file(file_path: str, file_content=None) -> Dict:
        """
        Parse CV file based on extension
        
        Args:
            file_path: Path or filename with extension
            file_content: File content (if already loaded)
            
        Returns:
            Dictionary with parsed CV data
        """
        
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == '.json':
            return CVParser.parse_json(file_content or file_path)
        elif extension == '.pdf':
            return CVParser.parse_pdf(file_path, file_content)
        elif extension in ['.docx', '.doc']:
            return CVParser.parse_docx(file_path, file_content)
        elif extension in ['.txt', '.text']:
            return CVParser.parse_txt(file_content or file_path)
        else:
            logger.warning(f"Unknown format: {extension}")
            return CVParser.parse_txt(file_content or file_path)
    
    @staticmethod
    def parse_json(file_content) -> Dict:
        """Parse JSON CV"""
        try:
            if isinstance(file_content, str):
                return json.loads(file_content)
            return file_content
        except Exception as e:
            logger.error(f"Error parsing JSON: {e}")
            return {}
    
    @staticmethod
    def parse_pdf(file_path: str, file_content=None) -> Dict:
        """Parse PDF CV using PyPDF2"""
        try:
            import PyPDF2
            import io
            
            if file_content:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            else:
                pdf_reader = PyPDF2.PdfReader(open(file_path, 'rb'))
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            
            return CVParser._extract_structured_data(text)
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return {}
    
    @staticmethod
    def parse_docx(file_path: str, file_content=None) -> Dict:
        """Parse DOCX CV"""
        try:
            from docx import Document
            import io
            
            if file_content:
                doc = Document(io.BytesIO(file_content))
            else:
                doc = Document(file_path)
            
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            return CVParser._extract_structured_data(text)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return {}
    
    @staticmethod
    def parse_txt(file_content: str) -> Dict:
        """Parse TXT CV"""
        try:
            if isinstance(file_content, bytes):
                text = file_content.decode('utf-8')
            else:
                text = file_content
            
            return CVParser._extract_structured_data(text)
            
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            return {}
    
    @staticmethod
    def _extract_structured_data(text: str) -> Dict:
        """
        Extract structured data from raw CV text
        Uses pattern matching to find common CV sections
        """
        
        lines = text.split('\n')
        
        # Initialize structure
        cv_data = {
            "name": "",
            "email": "",
            "phone": "",
            "degree": "",
            "university": "",
            "years_experience": 0,
            "current_role": "",
            "current_company": "",
            "experience": [],
            "skills": [],
            "soft_skills": [],
            "education": [],
            "raw_text": text[:500]  # First 500 chars for reference
        }
        
        # Extract name (usually first non-empty line)
        for line in lines[:5]:
            if line.strip() and len(line.strip()) < 50:
                cv_data["name"] = line.strip()
                break
        
        # Extract email
        for line in text.split('\n'):
            if '@' in line and '.' in line:
                email = line.strip()
                if email.count('@') == 1:
                    cv_data["email"] = email
                    break
        
        # Extract phone
        import re
        phone_pattern = r'(\+?1?\s*)?(\d{3}[-.]?\d{3}[-.]?\d{4})'
        for match in re.finditer(phone_pattern, text):
            cv_data["phone"] = match.group(0)
            break
        
        # Extract education keywords
        education_keywords = ['bs ', 'ba ', 'ms ', 'ma ', 'phd', 'mba', 'degree', 'certificate']
        for line in lines:
            line_lower = line.lower()
            for keyword in education_keywords:
                if keyword in line_lower:
                    cv_data["degree"] = line.strip()
                    break
        
        # Extract university keywords
        university_keywords = ['university', 'college', 'institute', 'school']
        for line in lines:
            line_lower = line.lower()
            for keyword in university_keywords:
                if keyword in line_lower:
                    cv_data["university"] = line.strip()
                    break
        
        # Extract experience (look for common job keywords)
        experience_keywords = ['engineer', 'developer', 'manager', 'director', 'lead', 'analyst', 'specialist']
        current_role = ""
        
        for line in lines:
            line_lower = line.lower()
            for keyword in experience_keywords:
                if keyword in line_lower and len(line.strip()) < 100:
                    if not current_role:
                        current_role = line.strip()
                        cv_data["current_role"] = current_role
                    
                    cv_data["experience"].append({
                        "role": line.strip(),
                        "company": "Unknown",
                        "duration": "N/A",
                        "description": ""
                    })
                    break
        
        # Extract skills (look for common skill keywords)
        skill_keywords = ['python', 'java', 'javascript', 'sql', 'aws', 'kubernetes', 'git',
                         'react', 'django', 'tensorflow', 'pytorch', 'machine learning', 'ai', 'ml']
        
        for line in lines:
            line_lower = line.lower()
            for skill in skill_keywords:
                if skill in line_lower:
                    if skill not in cv_data["skills"]:
                        cv_data["skills"].append(skill.title())
        
        # Extract soft skills
        soft_skills_keywords = ['leadership', 'communication', 'problem-solving', 'teamwork',
                              'collaboration', 'creativity', 'adaptability', 'mentoring', 'presentation']
        
        for line in lines:
            line_lower = line.lower()
            for skill in soft_skills_keywords:
                if skill in line_lower:
                    if skill not in cv_data["soft_skills"]:
                        cv_data["soft_skills"].append(skill.title())
        
        # Estimate years of experience from keywords
        for line in lines:
            line_lower = line.lower()
            if 'year' in line_lower:
                import re
                numbers = re.findall(r'\d+', line)
                if numbers:
                    cv_data["years_experience"] = max(cv_data["years_experience"], int(numbers[0]))
        
        return cv_data


def parse_cv_streamlit(uploaded_file) -> Dict:
    """
    Parse CV from Streamlit uploaded file
    Works with any file format
    """
    
    try:
        file_name = uploaded_file.name
        file_content = uploaded_file.read()
        
        # Try to parse based on file type
        extension = os.path.splitext(file_name)[1].lower()
        
        if extension == '.json':
            return CVParser.parse_json(json.loads(file_content.decode('utf-8')))
        elif extension == '.pdf':
            return CVParser.parse_pdf(None, file_content)
        elif extension in ['.docx', '.doc']:
            return CVParser.parse_docx(None, file_content)
        else:
            return CVParser.parse_txt(file_content)
            
    except Exception as e:
        logger.error(f"Error parsing CV: {e}")
        return {
            "name": "Error parsing file",
            "error": str(e),
            "raw_text": ""
        }
